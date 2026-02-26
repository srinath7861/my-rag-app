"""
Ingest PDF, DOCX, and URL into the RAG knowledge base (Chroma).
"""
import re
import uuid
from pathlib import Path

import requests
from bs4 import BeautifulSoup

from config import CHUNK_SIZE, CHUNK_OVERLAP, KNOWLEDGE_PATH, URL_CONTENT_PATH, QNA_PATH, DOCUMENTS_DIR
from rag_core import chunk_text, add_chunks_to_collection, delete_chunks_by_source

# Optional deps for PDF/DOCX
try:
    import pypdf
except ImportError:
    pypdf = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_text_pdf(path_or_bytes) -> str:
    """Extract text from a PDF file path or bytes."""
    if pypdf is None:
        raise ImportError("Install pypdf: pip install pypdf")
    if isinstance(path_or_bytes, (str, Path)):
        reader = pypdf.PdfReader(str(path_or_bytes))
    else:
        from io import BytesIO
        reader = pypdf.PdfReader(BytesIO(path_or_bytes))
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return _clean_text("\n".join(parts))


def extract_text_docx(path_or_bytes) -> str:
    """Extract text from a DOCX file path or bytes."""
    if DocxDocument is None:
        raise ImportError("Install python-docx: pip install python-docx")
    if isinstance(path_or_bytes, (str, Path)):
        doc = DocxDocument(str(path_or_bytes))
    else:
        from io import BytesIO
        doc = DocxDocument(BytesIO(path_or_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return _clean_text("\n".join(parts))


def extract_text_url(url: str, timeout: int = 15) -> str:
    """Fetch URL and extract main text (strip script/style, get body text)."""
    headers = {"User-Agent": "RAGBot/1.0 (Knowledge base ingestion)"}
    resp = requests.get(url, timeout=timeout, headers=headers)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return _clean_text(text)


def ingest_text(
    text: str,
    source_label: str,
    metadata_base: dict | None = None,
) -> int:
    """
    Chunk text and add to Chroma. source_label is used in 'source' metadata.
    Returns number of chunks added.
    """
    if not text or not text.strip():
        return 0
    chunks = chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)
    if not chunks:
        return 0
    base = metadata_base or {}
    metadatas = [{**base, "source": source_label} for _ in chunks]
    ids = [str(uuid.uuid4()) for _ in chunks]
    add_chunks_to_collection(chunks, metadatas=metadatas, ids=ids)
    return len(chunks)


def ingest_pdf(path_or_bytes, filename: str | None = None) -> int:
    """Ingest a PDF. filename used as source label (e.g. 'policy.pdf')."""
    text = extract_text_pdf(path_or_bytes)
    label = filename or (getattr(path_or_bytes, "name", None) if hasattr(path_or_bytes, "name") else "document.pdf")
    if isinstance(path_or_bytes, (str, Path)):
        label = label or str(Path(path_or_bytes).name)
    return ingest_text(text, source_label=label, metadata_base={"type": "pdf"})


def ingest_docx(path_or_bytes, filename: str | None = None) -> int:
    """Ingest a DOCX. filename used as source label."""
    text = extract_text_docx(path_or_bytes)
    label = filename or (getattr(path_or_bytes, "name", None) if hasattr(path_or_bytes, "name") else "document.docx")
    if isinstance(path_or_bytes, (str, Path)):
        label = label or str(Path(path_or_bytes).name)
    return ingest_text(text, source_label=label, metadata_base={"type": "docx"})


def _append_url_content_to_file(url: str, text: str) -> None:
    """Append extracted URL text to url_content.txt for persistent storage and RAG training."""
    path = Path(URL_CONTENT_PATH)
    if not path.is_absolute():
        path = Path(__file__).resolve().parent / path
    path.parent.mkdir(parents=True, exist_ok=True)
    separator = "\n\n--- URL: {} ---\n\n".format(url)
    with path.open("a", encoding="utf-8") as f:
        f.write(separator)
        f.write(text)
        f.write("\n")


def ingest_url(url: str) -> int:
    """Fetch URL, extract text, save to url_content.txt, and ingest into Chroma."""
    text = extract_text_url(url)
    _append_url_content_to_file(url, text)
    return ingest_text(text, source_label=url, metadata_base={"type": "url", "url": url})


def ingest_knowledge_file(path: str | None = None) -> int:
    """Ingest a plain-text knowledge file (e.g. knowledge.txt)."""
    path = Path(path or KNOWLEDGE_PATH)
    if not path.exists():
        raise FileNotFoundError(str(path))
    text = path.read_text(encoding="utf-8", errors="replace")
    return ingest_text(text, source_label=path.name, metadata_base={"type": "txt"})


def _url_content_path() -> Path:
    p = Path(URL_CONTENT_PATH)
    if not p.is_absolute():
        p = Path(__file__).resolve().parent / p
    return p


def parse_url_content_file() -> list[dict]:
    """Return list of {url, content} from url_content.txt."""
    path = _url_content_path()
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8", errors="replace")
    entries = []
    for block in text.split("--- URL:"):
        block = block.strip()
        if not block:
            continue
        first_newline = block.find("\n")
        if first_newline >= 0:
            url = block[:first_newline].strip()
            content = block[first_newline:].strip()
        else:
            url = block
            content = ""
        if url:
            entries.append({"url": url, "content": content})
    return entries


def rewrite_url_content_file(entries: list[dict]) -> None:
    """Write url_content.txt with given list of {url, content} (used after removing one URL)."""
    path = _url_content_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    parts = []
    for e in entries:
        parts.append("--- URL: {} ---\n\n{}".format(e["url"], e.get("content", "")))
    path.write_text("\n\n".join(parts), encoding="utf-8")


def remove_url_from_knowledge_base(url: str) -> None:
    """Remove URL's chunks from Chroma and remove from url_content.txt."""
    delete_chunks_by_source(url)
    entries = [e for e in parse_url_content_file() if e["url"] != url]
    rewrite_url_content_file(entries)


def update_url_content(url: str, new_content: str) -> None:
    """Replace URL's stored content and re-ingest (delete old chunks, add new)."""
    delete_chunks_by_source(url)
    entries = parse_url_content_file()
    for e in entries:
        if e["url"] == url:
            e["content"] = new_content
            break
    else:
        entries.append({"url": url, "content": new_content})
    rewrite_url_content_file(entries)
    if new_content.strip():
        ingest_text(new_content, source_label=url, metadata_base={"type": "url", "url": url})


def _documents_path() -> Path:
    p = Path(DOCUMENTS_DIR)
    if not p.is_absolute():
        p = Path(__file__).resolve().parent / p
    return p


def list_documents() -> list[dict]:
    """List named text documents (id = filename without ext)."""
    base = _documents_path()
    if not base.exists():
        return []
    out = []
    for f in base.iterdir():
        if f.suffix == ".txt" and f.is_file():
            out.append({"id": f.stem, "name": f.name})
    return sorted(out, key=lambda x: x["name"])


def get_document_content(doc_id: str) -> str:
    """Get content of a document by id (filename stem)."""
    base = _documents_path()
    base.mkdir(parents=True, exist_ok=True)
    f = base / (doc_id + ".txt")
    if not f.exists():
        return ""
    return f.read_text(encoding="utf-8", errors="replace")


def save_document(doc_id: str, name: str, content: str) -> int:
    """Save document and ingest. doc_id used as source in Chroma."""
    base = _documents_path()
    base.mkdir(parents=True, exist_ok=True)
    f = base / (doc_id + ".txt")
    f.write_text(content, encoding="utf-8")
    delete_chunks_by_source(doc_id)
    if content.strip():
        return ingest_text(content, source_label=doc_id, metadata_base={"type": "doc", "name": name})
    return 0


def delete_document(doc_id: str) -> None:
    """Remove document file and its chunks from Chroma."""
    delete_chunks_by_source(doc_id)
    base = _documents_path()
    f = base / (doc_id + ".txt")
    if f.exists():
        f.unlink()


def _qna_path() -> Path:
    p = Path(QNA_PATH)
    if not p.is_absolute():
        p = Path(__file__).resolve().parent / p
    return p


def parse_qna_file() -> list[dict]:
    """Return list of {question, answer} from qna.txt. Format: Q: ... A: ..."""
    path = _qna_path()
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8", errors="replace")
    entries = []
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        q, a = "", ""
        if block.startswith("Q:") or block.startswith("q:"):
            idx = block.find("\n")
            if idx >= 0:
                q = block[2:idx].strip()
                rest = block[idx:].strip()
                if rest.startswith("A:") or rest.startswith("a:"):
                    a = rest[2:].strip()
            else:
                q = block[2:].strip()
        entries.append({"question": q, "answer": a})
    return entries


def append_qna(question: str, answer: str) -> int:
    """Append Q&A to qna.txt and ingest as text 'Q: ... A: ...'."""
    path = _qna_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    line = "Q: {}\nA: {}\n\n".format(question.strip(), answer.strip())
    with path.open("a", encoding="utf-8") as f:
        f.write(line)
    text = "Q: {}\nA: {}".format(question.strip(), answer.strip())
    return ingest_text(text, source_label="qna", metadata_base={"type": "qna"})


def delete_all_qna() -> None:
    """Remove all Q&A chunks from Chroma and clear qna.txt."""
    delete_chunks_by_source("qna")
    path = _qna_path()
    if path.exists():
        path.write_text("", encoding="utf-8")


def _write_qna_file(entries: list[dict]) -> None:
    """Write qna.txt from list of {question, answer}."""
    path = _qna_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = ["Q: {}\nA: {}".format(e.get("question", "").strip(), e.get("answer", "").strip()) for e in entries]
    path.write_text("\n\n".join(lines) + ("\n\n" if lines else ""), encoding="utf-8")


def delete_qna_at_index(index: int) -> None:
    """Remove the Q&A pair at index (0-based), rewrite qna.txt, and re-ingest remaining Q&A."""
    entries = parse_qna_file()
    if index < 0 or index >= len(entries):
        raise IndexError("Q&A index out of range")
    entries.pop(index)
    _write_qna_file(entries)
    delete_chunks_by_source("qna")
    if entries:
        text = "\n\n".join("Q: {}\nA: {}".format(e.get("question", ""), e.get("answer", "")) for e in entries)
        ingest_text(text, source_label="qna", metadata_base={"type": "qna"})


def reingest_all_sources() -> int:
    """Clear Chroma and ingest knowledge.txt, all URLs, qna.txt, and all documents. Returns total chunks."""
    from rag_core import clear_collection
    clear_collection()
    n = 0
    path_k = Path(KNOWLEDGE_PATH)
    if not path_k.is_absolute():
        path_k = Path(__file__).resolve().parent / path_k
    if path_k.exists():
        n += ingest_knowledge_file(str(path_k))
    for e in parse_url_content_file():
        if e.get("content"):
            n += ingest_text(e["content"], source_label=e["url"], metadata_base={"type": "url", "url": e["url"]})
    path_q = _qna_path()
    if path_q.exists():
        text = path_q.read_text(encoding="utf-8", errors="replace")
        if text.strip():
            n += ingest_text(text, source_label="qna", metadata_base={"type": "qna"})
    for doc in list_documents():
        content = get_document_content(doc["id"])
        if content.strip():
            n += ingest_text(content, source_label=doc["id"], metadata_base={"type": "doc", "name": doc["name"]})
    return n
